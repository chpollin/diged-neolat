#!/usr/bin/env python3
"""
Lucina Edition.docx to TEI XML converter - Final Version
Fixes Praefatio headers, colophon handling, and lyric meter detection
"""

import re
import logging
from pathlib import Path
from docx import Document
from lxml import etree
from typing import Dict, List, Optional, Tuple
from collections import defaultdict

# Configure logging
logging.basicConfig(
   level=logging.INFO,
   format='%(levelname)s: %(message)s'
)
logger = logging.getLogger(__name__)


class LucinaToTEI:
   """Convert Lucina Edition.docx to TEI XML - Final corrected version"""
   
   def __init__(self, docx_path: str):
       """Initialize with document path"""
       self.doc_path = Path(docx_path)
       logger.info(f"Loading: {self.doc_path.name}")
       self.doc = Document(docx_path)
       
       # TEI namespace setup
       self.TEI_NS = "http://www.tei-c.org/ns/1.0"
       self.XML_NS = "http://www.w3.org/XML/1998/namespace"
       self.nsmap = {None: self.TEI_NS, 'xml': self.XML_NS}
       
       # State tracking
       self.current_book = None
       self.current_poem = None
       self.poems = []
       self.verse_buffer = []
       self.in_praefatio = False
       self.praefatio_lines_seen = 0  # Count actual lines after headers
       self.stats = defaultdict(int)
       self.poem_tracker = defaultdict(list)
       self.unprocessed_lines = []
       self.colophon_lines = []  # Store colophon separately
       
       logger.info(f"Total paragraphs: {len(self.doc.paragraphs)}")
   
   def process(self) -> etree.Element:
       """Main processing method"""
       logger.info("="*60)
       logger.info("Starting final document processing...")
       logger.info("Fixing Praefatio headers and colophon")
       logger.info("="*60)
       
       # Create TEI structure
       tei = etree.Element('TEI', nsmap=self.nsmap)
       text = etree.SubElement(tei, 'text')
       front = etree.SubElement(text, 'front')
       body = etree.SubElement(text, 'body')
       back = etree.SubElement(text, 'back')
       
       # Process all paragraphs
       for idx, para in enumerate(self.doc.paragraphs):
           if idx % 100 == 0 and idx > 0:
               logger.debug(f"Progress: {idx}/{len(self.doc.paragraphs)} paragraphs")
           self._process_paragraph(para, idx)
       
       # Save final poem if exists
       if self.current_poem and self.verse_buffer:
           self._save_current_poem()
       
       # Build TEI structure
       self._build_tei_structure(front, body, back)
       
       # Report results
       self._report_results()
       
       return tei
   
   def _process_paragraph(self, para, idx: int):
       """Process a single paragraph"""
       text = para.text.strip()
       if not text:
           return
       
       # Check for colophon FIRST (before other processing)
       if self._is_colophon(text):
           self._handle_colophon(text)
           return
       
       # Check for structural elements
       if self._is_book_header(text):
           self._handle_book_header(text)
       elif self._is_poem_header(text):
           self._handle_poem_header(text, idx)
       elif self.current_poem is not None:
           # We're in a poem, add as verse line
           self._handle_verse_line(text)
       elif self.in_praefatio:
           # We're in Praefatio after header
           self._handle_praefatio_content(text)
       else:
           # Track what we couldn't process
           self.unprocessed_lines.append((idx, text[:50]))
           self.stats['uncertain'] += 1
   
   def _is_colophon(self, text: str) -> bool:
       """Check if text is a colophon (at end of document)"""
       return 'Actum Papiae' in text or 'CDiis Immor' in text or 'Quarto Nonas Augustas' in text
   
   def _handle_colophon(self, text: str):
       """Handle colophon - store separately, don't add to poem"""
       self.colophon_lines.append(text)
       logger.info(f"Found colophon: {text[:50]}...")
       self.stats['colophon'] += 1
   
   def _is_book_header(self, text: str) -> bool:
       """Check if text is a book header"""
       if text == "Praefatio":
           return True
       if re.match(r'^Buch\s+[IVX]+$', text):
           return True
       if re.match(r'^B\[uch\]\s+[IVX]+$', text):
           return True
       return False
   
   def _is_poem_header(self, text: str) -> bool:
       """Check if text is a poem header - ALL patterns"""
       patterns = [
           # Standard Ad pattern
           r'^([IVX]+),\s*(\d+)\s+Ad\s+(.+?)\.?$',
           # In pattern (invectives/epitaphs)
           r'^([IVX]+),\s*(\d+)\s+In\s+(.+?)\.?$',
           # No preposition pattern (direct title/name)
           r'^([IVX]+),\s*(\d+)\s+([A-Z][a-z].+?)\.?$',
           # Numeric book numbers (if any)
           r'^(\d+),\s*(\d+)\s+(?:Ad|In)\s+(.+?)\.?$',
       ]
       return any(re.match(p, text) for p in patterns)
   
   def _handle_book_header(self, text: str):
       """Process book header"""
       # Save current poem if exists
       if self.current_poem and self.verse_buffer:
           self._save_current_poem()
       
       if text == "Praefatio":
           self.current_book = 'praefatio'
           self.in_praefatio = True
           self.praefatio_lines_seen = 0  # Reset counter
           self.stats['books'] += 1
           # Create praefatio poem structure
           self.current_poem = {
               'book': 'praefatio',
               'number': '0',
               'type': 'praefatio',
               'lines': []
           }
           self.verse_buffer = []
           logger.info("✓ Found: Praefatio")
           
       elif "Buch" in text or "B[uch]" in text:
           # End praefatio if we were in it
           if self.in_praefatio:
               self.in_praefatio = False
               if self.verse_buffer:
                   self._save_current_poem()
           
           match = re.search(r'[IVX]+', text)
           if match:
               self.current_book = match.group()
               self.stats['books'] += 1
               logger.info(f"✓ Found: Book {self.current_book}")
   
   def _handle_praefatio_content(self, text: str):
       """Handle content in Praefatio - skip the first two header lines"""
       # Skip the headers: "Aurelii..." and "ad mecoenatem..."
       if 'Aurelii' in text and 'Albrisii' in text and 'praefatio' in text:
           logger.debug(f"Skipping Praefatio header 1: {text[:50]}")
           return
       if text.lower().startswith('ad mecoenatem') or text.lower().startswith('ad mecoenatum'):
           logger.debug(f"Skipping Praefatio header 2: {text[:50]}")
           return
       
       # Now we have actual verse lines
       self.verse_buffer.append(text)
       self.praefatio_lines_seen += 1
       self.stats['lines'] += 1
       logger.debug(f"Praefatio line {self.praefatio_lines_seen}: {text[:30]}")
   
   def _handle_poem_header(self, text: str, idx: int):
       """Process poem header - handles all patterns"""
       # Save previous poem
       if self.current_poem and self.verse_buffer:
           self._save_current_poem()
       
       # End praefatio state if we hit a poem
       if self.in_praefatio:
           self.in_praefatio = False
       
       # Try all patterns with their prepositions
       patterns = [
           (r'^([IVX]+),\s*(\d+)\s+Ad\s+(.+?)\.?$', 'Ad'),
           (r'^([IVX]+),\s*(\d+)\s+In\s+(.+?)\.?$', 'In'),
           (r'^([IVX]+),\s*(\d+)\s+([A-Z][a-z].+?)\.?$', ''),
           (r'^(\d+),\s*(\d+)\s+(Ad|In)\s+(.+?)\.?$', 'extract'),
       ]
       
       for pattern, prep_type in patterns:
           match = re.match(pattern, text)
           if match:
               # Extract components
               book = match.group(1)
               
               # Convert numeric to Roman if needed
               if book.isdigit():
                   book_map = {'1': 'I', '2': 'II', '3': 'III'}
                   book = book_map.get(book, self.current_book)
               
               number = match.group(2)
               
               # Handle different pattern structures
               if prep_type == 'extract':
                   preposition = match.group(3)
                   dedicatee = match.group(4).strip()
               else:
                   dedicatee = match.group(3).strip() if len(match.groups()) >= 3 else ""
                   preposition = prep_type
               
               # Clean dedicatee - remove any trailing rubric text
               if 'Aurelius' in dedicatee or 'Albrisii' in dedicatee or 'incipit' in dedicatee:
                   for keyword in ['Aurelius', 'Albrisii', 'incipit']:
                       if keyword in dedicatee:
                           dedicatee = dedicatee.split(keyword)[0].strip()
                           break
               
               # Create poem structure
               self.current_poem = {
                   'book': book,
                   'number': number,
                   'dedicatee': dedicatee,
                   'preposition': preposition,
                   'lines': []
               }
               
               self.verse_buffer = []
               self.stats['poems'] += 1
               self.poem_tracker[book].append(number)
               
               prep_display = f"{preposition} " if preposition else ""
               logger.info(f"✓ Found poem: {book}.{number} {prep_display}{dedicatee[:30]}")
               break
   
   def _handle_verse_line(self, text: str):
       """Process verse line"""
       if not text:
           return
       
       # Skip rubric-like lines
       if 'incipit' in text.lower() and 'aurelii' in text.lower():
           logger.debug(f"Skipping rubric-like line: {text[:50]}")
           return
       
       self.verse_buffer.append(text)
       self.stats['lines'] += 1
   
   def _save_current_poem(self):
       """Save current poem with its lines"""
       if self.current_poem and self.verse_buffer:
           self.current_poem['lines'] = self.verse_buffer.copy()
           self.poems.append(self.current_poem.copy())
           self.verse_buffer = []
   
   def _build_tei_structure(self, front, body, back):
       """Build TEI XML structure"""
       logger.info("\nBuilding TEI structure...")
       
       # Separate praefatio from other poems
       books = defaultdict(list)
       praefatio_poem = None
       
       for poem in self.poems:
           if poem.get('type') == 'praefatio' or poem['book'] == 'praefatio':
               praefatio_poem = poem
           else:
               books[poem['book']].append(poem)
       
       # Add Praefatio to front matter
       if praefatio_poem:
           self._add_praefatio(front, praefatio_poem)
       
       # Add books to body
       for book_num in ['I', 'II', 'III']:
           if book_num in books:
               book_div = etree.SubElement(body, 'div')
               book_div.set('type', 'book')
               book_div.set('n', str(['I', 'II', 'III'].index(book_num) + 1))
               book_div.set('{%s}id' % self.XML_NS, f'book{book_num}')
               
               # Add book heading
               head = etree.SubElement(book_div, 'head')
               head.set('type', 'book')
               head.text = f'Liber {self._roman_to_word(book_num)}'
               
               # Sort and add poems
               sorted_poems = sorted(books[book_num], 
                                   key=lambda p: int(p['number']))
               for poem in sorted_poems:
                   self._add_poem_to_tei(book_div, poem)
               
               logger.info(f"  Added Book {book_num}: {len(sorted_poems)} poems")
       
       # Add colophon to back matter if exists
       if self.colophon_lines:
           self._add_colophon(back)
   
   def _roman_to_word(self, roman: str) -> str:
       """Convert Roman numeral to Latin word"""
       return {'I': 'Primus', 'II': 'Secundus', 'III': 'Tertius'}.get(roman, roman)
   
   def _add_praefatio(self, front, poem):
       """Add praefatio to front matter - with correct line numbering"""
       div = etree.SubElement(front, 'div')
       div.set('type', 'praefatio')
       div.set('{%s}id' % self.XML_NS, 'praefatio')
       
       # Simple headers without duplication
       head1 = etree.SubElement(div, 'head')
       head1.set('type', 'author')
       head1.text = 'Aurelii Laurentii Albrisii praefatio in Lucinam'
       
       head2 = etree.SubElement(div, 'head')
       head2.set('type', 'dedication')
       head2.text = 'ad mecoenatem Cichum Simonetam'
       
       # Add lines in elegiac couplets - these should be the actual verses only
       lines = poem.get('lines', [])
       for i in range(0, len(lines), 2):
           lg = etree.SubElement(div, 'lg')
           lg.set('type', 'elegiac')
           
           for j in range(2):
               if i + j < len(lines):
                   l = etree.SubElement(lg, 'l')
                   l.set('n', str(i + j + 1))  # Correct numbering from 1
                   l.text = lines[i + j]
       
       self.stats['praefatio_lines'] = len(lines)
   
   def _add_poem_to_tei(self, book_div, poem):
       """Add poem to TEI structure"""
       poem_div = etree.SubElement(book_div, 'div')
       poem_div.set('type', 'poem')
       poem_div.set('n', poem['number'])
       poem_div.set('{%s}id' % self.XML_NS, f"poem-{poem['book']}.{poem['number']}")
       
       # Add poem number
       head_num = etree.SubElement(poem_div, 'head')
       head_num.set('type', 'number')
       head_num.text = f"{poem['book']}, {poem['number']}"
       
       # Add dedication with appropriate preposition
       if poem.get('dedicatee'):
           head_ded = etree.SubElement(poem_div, 'head')
           head_ded.set('type', 'dedication')
           
           prep = poem.get('preposition', '')
           if prep:
               head_ded.text = f"{prep} {poem['dedicatee']}"
           else:
               head_ded.text = poem['dedicatee']
       
       # Add lines with appropriate meter structure
       lines = poem.get('lines', [])
       if lines:
           meter_type = self._detect_meter(poem, lines)
           
           if meter_type == 'elegiac':
               # Add lines in elegiac couplets
               for i in range(0, len(lines), 2):
                   lg = etree.SubElement(poem_div, 'lg')
                   lg.set('type', 'elegiac')
                   
                   for j in range(2):
                       if i + j < len(lines):
                           l = etree.SubElement(lg, 'l')
                           l.set('n', str(i + j + 1))
                           l.text = lines[i + j]
           
           elif meter_type == 'lyric':
               # Lyric poems - single lg without type specification
               lg = etree.SubElement(poem_div, 'lg')
               for i, line_text in enumerate(lines, 1):
                   l = etree.SubElement(lg, 'l')
                   l.set('n', str(i))
                   l.text = line_text
           
           else:
               # Default: single lg for unknown meters
               lg = etree.SubElement(poem_div, 'lg')
               for i, line_text in enumerate(lines, 1):
                   l = etree.SubElement(lg, 'l')
                   l.set('n', str(i))
                   l.text = line_text
   
   def _detect_meter(self, poem: Dict, lines: List[str]) -> str:
       """Detect meter type for a poem"""
       # Known lyric poems (odd line counts or special meters)
       lyric_poems = [
           ('I', '18'),   # 3 lines
           ('I', '43'),   # 84 lines (sapphic)
           ('II', '6'),   # 51 lines
           ('III', '1'),  # 29 lines
           ('III', '10'), # 39 lines
           ('III', '36'), # 29 lines (hexameter)
           ('III', '43'), # 55 lines
           ('III', '47'), # Multiple lines, lyric
       ]
       
       if (poem['book'], poem['number']) in lyric_poems:
           return 'lyric'
       
       # Poems with even line count are likely elegiac
       if len(lines) % 2 == 0:
           return 'elegiac'
       
       # Odd line count = likely lyric
       return 'lyric'
   
   def _add_colophon(self, back):
       """Add colophon to back matter"""
       div = etree.SubElement(back, 'div')
       div.set('type', 'colophon')
       
       for line in self.colophon_lines:
           p = etree.SubElement(div, 'p')
           p.text = line
       
       logger.info(f"Added colophon with {len(self.colophon_lines)} lines")
   
   def _report_results(self):
       """Report processing results"""
       logger.info("\n" + "="*60)
       logger.info("PROCESSING COMPLETE:")
       logger.info("="*60)
       
       # Include Praefatio in count
       total_poems = self.stats['poems']
       if self.stats.get('praefatio_lines') > 0:
           total_poems += 1
       
       logger.info(f"✓ Successfully processed:")
       logger.info(f"  - Books: {self.stats['books']}")
       logger.info(f"  - Poems: {total_poems} (including Praefatio)")
       logger.info(f"  - Lines: {self.stats['lines']}")
       logger.info(f"  - Colophon sections: {self.stats['colophon']}")
       
       if self.stats.get('praefatio_lines'):
           logger.info(f"  - Praefatio lines: {self.stats['praefatio_lines']}")
           if self.stats['praefatio_lines'] == 16:
               logger.info("    ✓ Praefatio has correct 16 lines!")
           else:
               logger.warning(f"    ⚠ Praefatio has {self.stats['praefatio_lines']} lines (expected 16)")
       
       # Report poems per book
       for book in ['I', 'II', 'III']:
           if book in self.poem_tracker:
               logger.info(f"\nBook {book}: {len(self.poem_tracker[book])} poems")
               numbers = sorted([int(n) for n in self.poem_tracker[book]])
               
               # Check for gaps
               if numbers:
                   expected = set(range(1, max(numbers) + 1))
                   found = set(numbers)
                   missing = expected - found
                   if missing:
                       logger.warning(f"  Missing numbers: {sorted(missing)}")
       
       if self.unprocessed_lines:
           logger.info(f"\n✗ Unprocessed lines: {len(self.unprocessed_lines)}")
           if len(self.unprocessed_lines) <= 10:
               for idx, text in self.unprocessed_lines[:5]:
                   logger.info(f"  Para {idx}: {text}...")
       
       logger.info("\n" + "="*60)
       
       # Final statistics
       expected_total = 128  # Corrected: Praef + I:43 + II:37 + III:47
       capture_rate = (total_poems / expected_total) * 100 if expected_total > 0 else 0
       logger.info(f"Capture rate: {total_poems}/{expected_total} poems ({capture_rate:.1f}%)")
   
   def save_tei(self, output_path: str):
       """Save TEI XML to file"""
       tei = self.process()
       tree = etree.ElementTree(tei)
       
       output = Path(output_path)
       tree.write(
           str(output),
           pretty_print=True,
           xml_declaration=True,
           encoding='UTF-8'
       )
       
       logger.info(f"\nSaved: {output.name}")
       return output


def main():
   """Main execution"""
   import sys
   
   if len(sys.argv) < 2:
       print("Usage: python docx-to-tei-xml.py Edition.docx [output.xml]")
       sys.exit(1)
   
   input_file = sys.argv[1]
   output_file = sys.argv[2] if len(sys.argv) > 2 else 'lucina_final.xml'
   
   try:
       converter = LucinaToTEI(input_file)
       converter.save_tei(output_file)
       
       print(f"\nConversion complete!")
       print(f"Output: {output_file}")
       print(f"Expected: 128 poems (Praef:1 + I:43 + II:37 + III:47)")
       print("\nFinal version with:")
       print("  - Praefatio headers properly skipped")
       print("  - Colophon in back matter")
       print("  - Lyric poems with correct lg structure")
       
   except Exception as e:
       logger.error(f"Conversion failed: {e}")
       import traceback
       traceback.print_exc()
       sys.exit(1)


if __name__ == '__main__':
   main()